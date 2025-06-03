import os
from docx import Document
from docx.shared import Inches

image_folder='.'
doc=Document()
for filename in sorted(os.listdir(image_folder)):
    if filename.endswith("_2022.png"):
        image_path = os.path.join(image_folder,filename)

        doc.add_picture(image_path, width=Inches(5))
        doc.add_paragraph(f"Figure: {filename}" , style="Caption")
        doc.add_paragraph

output_path="RF_timeseries_2022.docx"
doc.save(output_path)

print(f"word document saved as {output_path}")
